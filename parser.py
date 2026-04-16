"""从新品文案 .docx 解析指定 SKU 的产品数据。

文案 .docx 第一个表格的结构：
- 行号 = 字段（如第 3 行是【产品编号】）
- 列号 = 产品（第 0 列是字段标签，第 1..N 列每列一个产品）

按 SKU 在第 3 行定位列号，然后从 FIELD_ROW_MAPPING 指定的行号读取该列的单元格。
字段命名与油猴脚本 cellIndexMappings 保持一致，Playwright 填充逻辑可直接复用。
"""

from __future__ import annotations

from pathlib import Path
from typing import Dict

from docx import Document


# 行号 → 字段名映射，对齐油猴脚本 cellIndexMappings（行号 row,1）
FIELD_ROW_MAPPING: Dict[int, str] = {
    1: "product_title",
    2: "category",
    3: "number",
    4: "search_keyword",
    5: "page_url",
    6: "description_en",
    11: "seo_keyword_en",
    12: "seo_description_en",
    14: "price",
    17: "stock",
    18: "is_coming",
    19: "presale_price",
    20: "pre_discount",
    21: "length",
    22: "width",
    23: "height",
    24: "weight",
    29: "is_batteries",
}


class DocxParseError(Exception):
    """.docx 解析失败时抛出。"""


def parse_docx(docx_path: str | Path, sku: str) -> Dict[str, str]:
    """从多产品文案 .docx 中按 SKU 定位并读取该产品的 18 个字段。

    Args:
        docx_path: .docx 文件路径
        sku: 产品编号，对应表格第 3 行（索引）的单元格值

    Returns:
        字段名 → 值 的字典，缺失字段填空字符串

    Raises:
        DocxParseError: 文件不存在、无表格、或 SKU 未在第 3 行找到
    """
    path = Path(docx_path).expanduser().resolve()
    if not path.exists():
        raise DocxParseError(f"文件不存在: {path}")
    if path.suffix.lower() != ".docx":
        raise DocxParseError(f"不是 .docx 文件: {path}")

    doc = Document(str(path))
    if not doc.tables:
        raise DocxParseError("文档中未找到表格")

    table = doc.tables[0]
    if len(table.rows) < 4:
        raise DocxParseError(f"表格行数不足（预期 ≥30，实际 {len(table.rows)}）")

    # 第 3 行（产品编号）逐列匹配 SKU
    sku_row = table.rows[3]
    col_idx = None
    for j, cell in enumerate(sku_row.cells):
        if j == 0:
            continue
        if cell.text.strip() == sku:
            col_idx = j
            break

    if col_idx is None:
        all_skus = [
            sku_row.cells[j].text.strip()
            for j in range(1, len(sku_row.cells))
            if sku_row.cells[j].text.strip()
        ]
        raise DocxParseError(
            f"文案第 3 行未找到 SKU={sku!r}；该行实际值：{all_skus}"
        )

    data: Dict[str, str] = {}
    total_rows = len(table.rows)
    for row_idx, field_name in FIELD_ROW_MAPPING.items():
        if row_idx >= total_rows:
            data[field_name] = ""
            continue
        row = table.rows[row_idx]
        if col_idx >= len(row.cells):
            data[field_name] = ""
            continue
        data[field_name] = row.cells[col_idx].text.strip()

    return data


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) < 3:
        print("用法: python parser.py <docx路径> <SKU>")
        sys.exit(1)
    result = parse_docx(sys.argv[1], sys.argv[2])
    print(json.dumps(result, ensure_ascii=False, indent=2))
